import streamlit as st
import google.genai as genai
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
import json
import os
import io
import time
from datetime import datetime
from pypdf import PdfReader
from docx import Document

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="AI Çeviri Asistanı", page_icon="🧠", layout="wide")

# --- SABİTLER ---
SCOPES = ['https://www.googleapis.com/auth/drive']
REDIRECT_URI = "urn:ietf:wg:oauth:2.0:oob"
ANA_KLASOR_ADI = "CEVIRI_PROJELERI_V2"
TOKEN_FILE = "token.json" # Kalıcı giriş için anahtar dosyası

# --- 1. GÜVENLİK VE GİRİŞ ---
def check_app_password():
    """Basit uygulama şifresi kontrolü."""
    if "auth_success" not in st.session_state:
        st.session_state.auth_success = False

    if not st.session_state.auth_success:
        st.markdown("## 🔒 Güvenlik Kilidi")
        pwd = st.text_input("Uygulama Şifresi:", type="password")
        if st.button("Giriş"):
            # Şifreyi buraya kendin belirle (Örn: 1234)
            if pwd == "1234": 
                st.session_state.auth_success = True
                st.rerun()
            else:
                st.error("Yanlış şifre!")
        st.stop()

def get_google_creds():
    """Token dosyasından yetkiyi okur, yoksa login ister."""
    creds = None
    # 1. Kayıtlı token var mı?
    if os.path.exists(TOKEN_FILE):
        try:
            creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        except:
            os.remove(TOKEN_FILE) # Bozuksa sil
            
    # 2. Token geçerli mi?
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                from google.auth.transport.requests import Request
                creds.refresh(Request())
                # Yenilenen tokenı kaydet
                with open(TOKEN_FILE, 'w') as token:
                    token.write(creds.to_json())
            except:
                creds = None # Yenilenemedi, sıfırdan al

    # 3. Hala yetki yoksa OAuth başlat
    if not creds:
        if "oauth" not in st.secrets:
            st.error("Secrets ayarı eksik!")
            st.stop()
            
        client_config = json.loads(st.secrets["oauth"]["CLIENT_CONFIG"])
        flow = Flow.from_client_config(client_config, scopes=SCOPES, redirect_uri=REDIRECT_URI)
        
        st.title("Google ile Bağlan (Tek Seferlik)")
        auth_url, _ = flow.authorization_url(prompt='consent')
        st.markdown(f"1. [İzin Linkine Tıkla]({auth_url})")
        code = st.text_input("2. Kodu Yapıştır:")
        
        if code:
            flow.fetch_token(code=code)
            creds = flow.credentials
            # Token'ı dosyaya kaydet (Kalıcılık sağlar!)
            with open(TOKEN_FILE, 'w') as token:
                token.write(creds.to_json())
            st.rerun()
        st.stop()
        
    return creds

# --- 2. DRIVE DOSYA YÖNETİMİ ---
def get_drive_service(creds):
    return build('drive', 'v3', credentials=creds)

def get_or_create_folder(service, folder_name, parent_id=None):
    q = f"name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    if parent_id: q += f" and '{parent_id}' in parents"
    
    results = service.files().list(q=q, fields="files(id)").execute()
    items = results.get('files', [])
    
    if items: return items[0]['id']
    
    metadata = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
    if parent_id: metadata['parents'] = [parent_id]
    
    folder = service.files().create(body=metadata, fields='id').execute()
    return folder.get('id')

def upload_file_content(service, folder_id, filename, content, mime_type):
    """Metin veya Binary içeriği dosyaya yazar/günceller."""
    # Dosya var mı?
    q = f"name = '{filename}' and '{folder_id}' in parents and trashed = false"
    results = service.files().list(q=q, fields="files(id)").execute()
    items = results.get('files', [])

    if isinstance(content, str): content = content.encode('utf-8')
    media = MediaIoBaseUpload(io.BytesIO(content), mimetype=mime_type, resumable=True)

    if items:
        # Güncelle
        service.files().update(fileId=items[0]['id'], media_body=media).execute()
        return items[0]['id']
    else:
        # Yarat
        meta = {'name': filename, 'parents': [folder_id]}
        f = service.files().create(body=meta, media_body=media, fields='id').execute()
        return f.get('id')

def read_file_content(service, folder_id, filename):
    """Dosya içeriğini okur (txt/json)."""
    q = f"name = '{filename}' and '{folder_id}' in parents and trashed = false"
    results = service.files().list(q=q, fields="files(id)").execute()
    items = results.get('files', [])
    
    if not items: return ""
    
    request = service.files().get_media(fileId=items[0]['id'])
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False: _, done = downloader.next_chunk()
    fh.seek(0)
    return fh.read().decode('utf-8')

# --- 3. YARDIMCI İŞLEMLER ---
def metni_parcala(metin):
    return [p.strip() for p in metin.split('\n\n') if p.strip()]

def word_olustur(paragraflar):
    doc = Document()
    for p in paragraflar:
        if p['durum'] == 'onaylandi': doc.add_paragraph(p['ceviri'])
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def ceviri_yap_gemini(metin, api_key, talimatlar, hafiza):
    try:
        client = genai.Client(api_key=api_key)
        prompt = f"""
        GÖREV: Aşağıdaki metni çevir.
        
        SİSTEM TALİMATLARI:
        {talimatlar}
        
        PROJE HAFIZASI (Öğrendiklerim):
        {hafiza}
        
        METİN:
        {metin}
        
        Sadece çeviriyi ver.
        """
        response = client.models.generate_content(model="gemini-2.5-pro", contents=prompt)
        return response.text.strip()
    except Exception as e: return f"Hata: {str(e)}"

# --- 4. UYGULAMA AKIŞI ---
check_app_password() # Önce şifre sor
creds = get_google_creds() # Sonra Google (Token varsa sormaz)
srv = get_drive_service(creds)
ana_id = get_or_create_folder(srv, ANA_KLASOR_ADI)

# --- SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Kontrol Paneli")
    api_key = st.text_input("Gemini API Key", type="password")
    if st.button("Projeleri Listele"):
        st.session_state.aktif_proje_id = None
        st.rerun()
    st.divider()
    if st.button("🔒 Güvenli Çıkış"):
        if os.path.exists(TOKEN_FILE): os.remove(TOKEN_FILE) # Token'ı sil
        st.session_state.auth_success = False
        st.rerun()

if "aktif_proje_id" not in st.session_state:
    st.session_state.aktif_proje_id = None

# --- EKRAN 1: PROJE LİSTESİ ---
if st.session_state.aktif_proje_id is None:
    st.title("📂 Projelerim")
    
    tabs = st.tabs(["Mevcut Projeler", "Yeni Proje Başlat"])
    
    with tabs[0]:
        q = f"'{ana_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        results = srv.files().list(q=q, fields="files(id, name)").execute()
        folders = results.get('files', [])
        
        if not folders: st.info("Henüz proje yok.")
        
        for f in folders:
            c1, c2 = st.columns([5,1])
            if c1.button(f"📁 {f['name']}", key=f['id']):
                st.session_state.aktif_proje_id = f['id']
                st.session_state.aktif_proje_adi = f['name']
                st.rerun()
                
            if c2.button("🗑️", key=f"d{f['id']}"):
                srv.files().delete(fileId=f['id']).execute()
                time.sleep(1); st.rerun()

    with tabs[1]:
        ad = st.text_input("Proje Adı")
        dosya = st.file_uploader("Dosya Yükle", type=['txt','docx','pdf'])
        talimat_giris = st.text_area("Bu Proje İçin Çeviri Talimatları:", 
                                     "Sen edebi bir çevirmensin. Tonu koru.")
        
        if st.button("Projeyi Oluştur") and ad and dosya:
            with st.spinner("Dosya sistemi ve veritabanı kuruluyor..."):
                # 1. Proje Klasörü
                p_id = get_or_create_folder(srv, ad, ana_id)
                
                # 2. Metni Oku
                def read_txt(f):
                    if f.name.endswith('.pdf'): return "".join([p.extract_text() for p in PdfReader(f).pages])
                    elif f.name.endswith('.docx'): return "\n\n".join([p.text for p in Document(f).paragraphs])
                    else: return f.read().decode('utf-8')
                
                ham_metin = read_txt(dosya)
                
                # 3. Dosyaları Drive'a At
                # Orijinal Dosya
                dosya.seek(0)
                upload_file_content(srv, p_id, f"ORIJINAL_{dosya.name}", dosya.read(), dosya.type)
                
                # Talimat Dosyası
                upload_file_content(srv, p_id, "TALIMATLAR.txt", talimat_giris, "text/plain")
                
                # Öğrendiklerim (Boş)
                upload_file_content(srv, p_id, "OGRENDIKLERIM.txt", "Henüz bir şey öğrenilmedi.", "text/plain")
                
                # Veritabanı (JSON)
                db_data = {
                    "meta": {"ad": ad, "tarih": str(datetime.now())},
                    "paragraflar": [{"id": i, "orjinal": p, "ceviri": "", "durum": "bekliyor"} 
                                    for i, p in enumerate(metni_parcala(ham_metin))]
                }
                upload_file_content(srv, p_id, "veritabani.json", json.dumps(db_data), "application/json")
                
                st.success("Proje Hazır!")
                time.sleep(1); st.rerun()

# --- EKRAN 2: PROJE ÇALIŞMA MASASI ---
else:
    pid = st.session_state.aktif_proje_id
    pname = st.session_state.aktif_proje_adi
    st.header(f"🛠️ {pname}")
    
    # Verileri Drive'dan Canlı Çek
    try:
        db_content = read_file_content(srv, pid, "veritabani.json")
        proje = json.loads(db_content) if db_content else None
        talimatlar = read_file_content(srv, pid, "TALIMATLAR.txt")
        hafiza = read_file_content(srv, pid, "OGRENDIKLERIM.txt")
    except:
        st.error("Veri okunamadı."); st.stop()
        
    paragraflar = proje["paragraflar"]
    if "cursor" not in st.session_state: st.session_state.cursor = 0
    
    # --- ÜST MENÜ (HAFIZA YÖNETİMİ) ---
    with st.expander("🧠 Yapay Zeka Hafızası & Talimatlar (Düzenle)"):
        c1, c2 = st.columns(2)
        yeni_talimat = c1.text_area("Talimatlar", talimatlar, height=150)
        yeni_hafiza = c2.text_area("Öğrendiklerim (Memory)", hafiza, height=150, help="Botun unutmamasını istediğin terimleri buraya ekle.")
        
        if st.button("Hafızayı Güncelle"):
            upload_file_content(srv, pid, "TALIMATLAR.txt", yeni_talimat, "text/plain")
            upload_file_content(srv, pid, "OGRENDIKLERIM.txt", yeni_hafiza, "text/plain")
            st.success("Hafıza güncellendi!")
            time.sleep(0.5); st.rerun()

    st.divider()

    # --- EDİTÖR ---
    idx = st.session_state.cursor
    # Navigasyon
    col_n1, col_n2, col_n3 = st.columns([1,1,5])
    if col_n1.button("⬅️"): st.session_state.cursor = max(0, idx-1); st.rerun()
    if col_n2.button("➡️"): st.session_state.cursor = min(len(paragraflar)-1, idx+1); st.rerun()
    st.caption(f"Paragraf: {idx+1} / {len(paragraflar)}")
    
    p = paragraflar[idx]
    
    c_sol, c_sag = st.columns(2)
    c_sol.info(p['orjinal'])
    
    with c_sag:
        if not p['ceviri'] and api_key and st.button("🤖 Çevir (Hafızalı)"):
            with st.spinner("Hafıza taranıyor ve çevriliyor..."):
                # Talimat + Hafıza + Metin gönderiliyor
                p['ceviri'] = ceviri_yap_gemini(p['orjinal'], api_key, talimatlar, hafiza)
                st.rerun()
        
        yeni_metin = st.text_area("Çeviri", p['ceviri'], height=200)
        
        if st.button("✅ Kaydet ve Dosyaları Güncelle", type="primary"):
            p['ceviri'] = yeni_metin
            p['durum'] = "onaylandi"
            
            # 1. Veritabanını Güncelle
            upload_file_content(srv, pid, "veritabani.json", json.dumps(proje), "application/json")
            
            # 2. Word Çıktısını Güncelle (Ceviri_Taslagi.docx)
            word_bytes = word_olustur(paragraflar)
            upload_file_content(srv, pid, f"CEVIRI_{pname}.docx", word_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # İlerle
            if idx < len(paragraflar)-1: st.session_state.cursor += 1
            st.toast("Kaydedildi! Word dosyası güncellendi.")
            time.sleep(0.5); st.rerun()
